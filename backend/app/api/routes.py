"""API route definitions."""
import io

import pdfplumber
from fastapi import APIRouter, HTTPException, UploadFile, File

from app.jobs.ranker import rank_jobs_for_resume, bulk_match_resumes
from app.jobs.repository import fetch_jobs_search
from app.models.schemas import (
    AnalyzeRequest, AnalyzeResponse,
    RankJobsRequest, RankJobsResponse,
    JobSearchRequest, JobSearchResponse, JobCard,
    BulkMatchRequest, BulkMatchResponse,
)
from app.nlp.matcher import analyze_match

router = APIRouter()


def _extract_text_from_pdf(contents: bytes) -> tuple[str, int]:
    """Extract plain text from a PDF byte stream.

    Returns (text, page_count).
    Raises ValueError if no text could be extracted.
    """
    with pdfplumber.open(io.BytesIO(contents)) as pdf:
        text = "\n".join(page.extract_text() or "" for page in pdf.pages)
        page_count = len(pdf.pages)
    text = text.strip()
    if not text:
        raise ValueError("Could not extract text from PDF. The file may be image-based.")
    return text, page_count


def _extract_text_from_docx(contents: bytes) -> tuple[str, int]:
    """Extract plain text from a DOCX byte stream.

    Returns (text, paragraph_count).
    Raises ValueError if no text could be extracted.
    """
    try:
        from docx import Document  # python-docx
    except ImportError as exc:
        raise ValueError("python-docx is not installed. Run: pip install python-docx") from exc

    doc = Document(io.BytesIO(contents))
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]

    # Also extract text from tables (common in resume templates)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                cell_text = cell.text.strip()
                if cell_text and cell_text not in paragraphs:
                    paragraphs.append(cell_text)

    text = "\n".join(paragraphs).strip()
    if not text:
        raise ValueError("Could not extract text from DOCX. The file may be empty or image-based.")
    return text, len(paragraphs)


@router.get("/health")
async def health_check():
    return {"status": "ok"}


@router.post("/upload-resume")
async def upload_resume(file: UploadFile = File(...)):
    """Extract text from an uploaded PDF or DOCX resume."""
    filename_lower = (file.filename or "").lower()

    if not (filename_lower.endswith(".pdf") or filename_lower.endswith(".docx")):
        raise HTTPException(
            status_code=422,
            detail="Only PDF and DOCX files are supported.",
        )

    try:
        contents = await file.read()

        if filename_lower.endswith(".pdf"):
            text, page_count = _extract_text_from_pdf(contents)
            return {"text": text, "filename": file.filename, "pages": page_count}

        else:  # .docx
            text, para_count = _extract_text_from_docx(contents)
            return {"text": text, "filename": file.filename, "pages": para_count}

    except (ValueError, HTTPException) as exc:
        raise HTTPException(status_code=422, detail=str(exc)) from exc
    except Exception as exc:
        raise HTTPException(status_code=500, detail=f"Failed to process file: {str(exc)}") from exc


@router.post("/analyze", response_model=AnalyzeResponse)
async def analyze(request: AnalyzeRequest):
    """Analyze resume against job description using NLP pipeline."""
    try:
        result = analyze_match(request.resume_text, request.job_description)
        return result
    except Exception as e:
        raise HTTPException(
            status_code=500,
            detail=f"NLP pipeline failed: {str(e)}",
        )


@router.post("/rank-jobs", response_model=RankJobsResponse)
async def rank_jobs(request: RankJobsRequest):
    """Rank stored jobs against a resume using the NLP matcher."""
    try:
        return rank_jobs_for_resume(request.model_dump())
    except (FileNotFoundError, RuntimeError) as exc:
        raise HTTPException(status_code=503, detail=str(exc)) from exc
    except Exception as exc:
        raise HTTPException(status_code=500, detail=f"Job ranking failed: {str(exc)}") from exc


@router.post("/jobs/search", response_model=JobSearchResponse)
async def search_jobs(request: JobSearchRequest):
    """Search and filter jobs from the corpus."""
    try:
        total, rows = fetch_jobs_search(
            keyword=request.keyword,
            province=request.province,
            city=request.city,
            broad_category=request.broad_category,
            page=request.page,
            page_size=request.page_size,
        )
        results = [
            JobCard(
                id=row["id"],
                job_title=row["job_title"],
                employer_name=row["employer_name"],
                city=row["city"],
                province=row["province"],
                salary=row["salary"],
                broad_category=row["broad_category"],
                date_posted=row["date_posted"],
                jd_preview=row["jd_preview"],
            )
            for row in rows
        ]
        return JobSearchResponse(total=total, page=request.page, page_size=request.page_size, results=results)
    except FileNotFoundError as exc:
        raise HTTPException(status_code=503, detail=str(exc)) from exc
    except Exception as exc:
        raise HTTPException(status_code=500, detail=f"Job search failed: {str(exc)}") from exc


@router.post("/bulk-match", response_model=BulkMatchResponse)
async def bulk_match(request: BulkMatchRequest):
    """Match multiple resumes against a filtered job set and return a leaderboard matrix."""
    try:
        result = bulk_match_resumes(request.model_dump())
        return result
    except FileNotFoundError as exc:
        raise HTTPException(status_code=503, detail=str(exc)) from exc
    except Exception as exc:
        raise HTTPException(status_code=500, detail=f"Bulk match failed: {str(exc)}") from exc
